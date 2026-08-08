"""多格式文档解析（RAG 入库入口）。

真实 RAG 系统的第一步不是「分块」，而是把异构格式的源文档统一解析成
可被分块的纯文本段落——PDF / DOCX / Markdown / 纯文本各有各的坑：
- PDF：多栏、页眉页脚、表格、扫描件需 OCR，朴素提取只能拿到底层文本流；
- DOCX：段落 + 表格 + 样式，需专用 SDK 抽取；
- Markdown：天然带标题层级，可按 heading 做结构感知切分；
- 纯文本：直接读。

本模块把上述格式归一为 ParsedSection 列表（每段带来源 meta），再交给
pipeline 做分块 / 嵌入 / 召回。

解析策略（优先 MinerU，轻量兜底）：
- 版面型格式（PDF / DOCX / PPTX / XLSX / 图片）优先走 MinerU 完成版面还原
  与结构化切片（标题 / 段落 / 表格 / 公式还原），输出 Markdown 后复用本模块
  的 heading 切分，得到结构感知的 section；
- MinerU 未安装或解析失败时，自动回退到内置轻量解析（pypdf / python-docx /
  纯文本），保证入库链路不崩溃；
- Markdown / 纯文本无需版面还原，直接走轻量解析。

依赖按格式懒加载，缺某格式依赖时仅在该格式解析时报错，不影响其他格式与导入。
MinerU 为可选依赖（见 requirements.txt 注释），不安装也能正常入库文本类文档。
"""

from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

logger = logging.getLogger(__name__)

_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")


@dataclass
class ParsedSection:
    """解析后的一个文本片段（对应后续一个分块单元的候选）。"""

    text: str
    meta: dict[str, Any] = field(default_factory=dict)


def _read_text(path: Path) -> list[ParsedSection]:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    raw = raw.strip()
    if not raw:
        return []
    return [ParsedSection(text=raw, meta={"format": "txt"})]


def _heading_level(h: str) -> int:
    return len(h.split(" ", 1)[0])


def _split_markdown(lines: list[str], source: str) -> list[ParsedSection]:
    """按标题层级切分，保留 heading 路径，使分块具备结构感知能力。

    每个 section 的 heading_path 包含其所属（含当前）标题，便于检索时
    还原文档结构上下文。MinerU 输出的 Markdown 也走同一套切分。
    """
    sections: list[ParsedSection] = []
    heading_stack: list[str] = []
    buf: list[str] = []

    def flush() -> None:
        nonlocal buf
        text = "\n".join(buf).strip()
        if text:
            sections.append(
                ParsedSection(
                    text=text,
                    meta={"format": "md", "source": source, "heading_path": list(heading_stack)},
                )
            )
        buf = []

    for line in lines:
        stripped = line.strip()
        m = _MD_HEADING.match(stripped)
        if m:
            # 先flush上一个 section（归属变更前的标题栈）
            flush()
            level = len(m.group(1))
            title = m.group(2).strip()
            while heading_stack and _heading_level(heading_stack[-1]) >= level:
                heading_stack.pop()
            heading_stack.append(f"{'#' * level} {title}")
            buf = [stripped]  # 标题本身作为新 section 起始上下文
        else:
            buf.append(line)
    flush()
    if not sections:  # 纯文本 markdown，无标题
        full = "\n".join(lines).strip()
        if full:
            sections.append(ParsedSection(text=full, meta={"format": "md", "source": source}))
    return sections


def _parse_markdown(path: Path) -> list[ParsedSection]:
    lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    return _split_markdown(lines, path.name)


def _parse_docx(path: Path) -> list[ParsedSection]:
    from docx import Document  # 懒加载

    doc = Document(str(path))
    sections: list[ParsedSection] = []
    base = path.name

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            sections.append(
                ParsedSection(
                    text=text,
                    meta={"format": "docx", "source": base, "kind": "paragraph"},
                )
            )

    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(" | ".join(cells))
        ttext = "\n".join(rows).strip()
        if ttext:
            sections.append(
                ParsedSection(
                    text=ttext,
                    meta={"format": "docx", "source": base, "kind": "table", "table_idx": ti},
                )
            )
    return sections


def _parse_pdf(path: Path) -> list[ParsedSection]:
    from pypdf import PdfReader  # 懒加载

    reader = PdfReader(str(path))
    sections: list[ParsedSection] = []
    base = path.name
    for pi, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            sections.append(
                ParsedSection(
                    text=text,
                    meta={"format": "pdf", "source": base, "page": pi + 1},
                )
            )
    return sections


# 版面型格式优先走 MinerU（版面还原 + 结构化切片）。
_MINERU_FORMATS = {".pdf", ".docx", ".pptx", ".xlsx", ".png", ".jpg", ".jpeg"}

_DISPATCH = {
    ".txt": _read_text,
    ".md": _parse_markdown,
    ".markdown": _parse_markdown,
    ".docx": _parse_docx,
    ".pdf": _parse_pdf,
}


def _try_mineru(path: Path) -> str | None:
    """尝试用 MinerU 把文档解析为 Markdown；任何失败都返回 None 走兜底。

    兼容 mineru.parse 的不同返回形态：同步对象（.markdown / to_dict）、
    纯字符串、dict，以及异步协程（在无事件循环的线程中 await 之）。
    """
    try:
        import mineru
    except Exception:
        return None
    try:
        result = mineru.parse(str(path))
        if hasattr(result, "__await__"):
            import asyncio

            try:
                asyncio.get_running_loop()
                # 已在事件循环内：避免嵌套事件循环，直接回退轻量解析。
                return None
            except RuntimeError:
                result = asyncio.run(result)
        if isinstance(result, str):
            md = result
        elif isinstance(result, dict):
            md = result.get("markdown")
        else:
            md = getattr(result, "markdown", None)
            if md is None and hasattr(result, "to_dict"):
                md = (result.to_dict() or {}).get("markdown")
        if not md:
            return None
        return str(md)
    except Exception as e:  # 未下载模型 / 解析异常等，均回退轻量解析
        logger.warning("MinerU 解析失败，回退轻量解析: %s", e)
        return None


def parse_document(path: str | Path) -> list[ParsedSection]:
    """解析单个文档为多段 ParsedSection。不支持的格式抛 ValueError。

    版面型格式优先 MinerU（版面还原 + 结构化切片）；其余或 MinerU 不可用
    时回退内置轻量解析。
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"文档不存在: {p}")
    ext = p.suffix.lower()

    # 1) 版面型格式优先 MinerU
    if ext in _MINERU_FORMATS:
        md = _try_mineru(p)
        if md is not None:
            return _split_markdown(md.splitlines(), p.name)

    # 2) 轻量兜底解析
    handler = _DISPATCH.get(ext)
    if handler is None:
        raise ValueError(
            f"不支持的文档格式: {ext}（支持 {sorted(_DISPATCH)}）"
        )
    try:
        sections = handler(p)
    except ImportError as e:
        raise ImportError(
            f"解析 {ext} 缺少依赖：{e}。请安装对应库（docx→python-docx，pdf→pypdf）。"
        ) from e
    # 统一补充来源文件名
    for s in sections:
        s.meta.setdefault("source", p.name)
    return sections


def parse_documents(paths: list[str | Path]) -> list[ParsedSection]:
    """批量解析，按顺序拼接所有 section。"""
    out: list[ParsedSection] = []
    for p in paths:
        out.extend(parse_document(p))
    return out


def sections_to_ragdocs(
    sections: list[ParsedSection], *, doc_id_prefix: str = "doc"
) -> list:
    """将解析后的段落映射为 pipeline 的 RagDoc，自动分配 chunk 友好的 id。"""
    from agent.rag.types import RagDoc

    docs: list[RagDoc] = []
    for i, s in enumerate(sections):
        docs.append(
            RagDoc(id=f"{doc_id_prefix}-{i}", text=s.text, meta=s.meta)
        )
    return docs
